"""Genera Modelo_de_Datos.docx — explicación pedagógica del modelo de datos
para audiencia sin conocimiento previo del tema (profesores, jurado)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "Modelo_de_Datos.docx"

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
for h, sz in [("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 13)]:
    s = styles[h]
    s.font.name = "Calibri"
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def shade(elem, color_hex):
    pPr = elem._p.get_or_add_pPr() if hasattr(elem, "_p") else elem._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def P(text="", *, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
    return p


def H(text, level):
    return doc.add_heading(text, level=level)


def CODE(code):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0)
        pf.left_indent = Inches(0.2)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        shade(p, "F2F2F2")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)


def BULLETS(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(it)


def NUMBERED(items):
    for it in items:
        p = doc.add_paragraph(style="List Number"); p.add_run(it)


def CALLOUT(text, color="FFF4CE"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(10)
    shade(p, color)


def TABLE(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, v in enumerate(row):
            c = t.rows[ri].cells[ci]; c.text = ""
            r = c.paragraphs[0].add_run(str(v)); r.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w


# =============================================================================
# TÍTULO
# =============================================================================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Modelo de Datos — Feedback de SAGE-EVA")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = st.add_run("Prueba de concepto + explicación didáctica")
sr.italic = True; sr.font.size = Pt(13); sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# =============================================================================
# 0. INTRODUCCIÓN
# =============================================================================
H("0. Introducción y propósito de este documento", 1)
P(
    "Este documento explica, desde cero, el modelo de datos que diseñé para almacenar "
    "el feedback que el asistente conversacional EVA recolecta durante un evento. Está "
    "escrito asumiendo que el lector no necesita tener conocimientos previos de bases "
    "de datos: cada concepto técnico se introduce con una definición y un ejemplo antes "
    "de usarse."
)
P(
    "El documento acompaña una prueba de concepto ejecutable (carpeta poc_data_model/) "
    "que crea la base, la puebla con datos representativos y corre queries analíticos. "
    "Cualquier afirmación que se hace acá se puede verificar corriendo los scripts."
)

# =============================================================================
# 1. EL PROBLEMA
# =============================================================================
H("1. El problema que resuelve este modelo", 1)

H("1.1 Contexto del proyecto", 2)
P(
    "SAGE es un sistema de gestión de eventos. Uno de sus módulos es EVA, un asistente "
    "conversacional que recibe a los invitados, les da indicaciones y recopila sus "
    "comentarios. EVA usa un modelo de lenguaje local (Ollama con llama3.2:3b) detrás "
    "de una API hecha en FastAPI."
)
P(
    "Hoy el sistema guarda el feedback en un archivo JSON plano. El archivo se reescribe "
    "entero cada vez que llega un comentario nuevo. La clasificación de cada comentario "
    "se hace con una heurística basada en palabras clave: si el mensaje del invitado "
    "contiene 'bueno', 'malo', 'organización', etc., se lo etiqueta automáticamente."
)

H("1.2 Limitaciones del enfoque actual", 2)
P("Esa solución es suficiente para un demo pero rompe contra cuatro problemas concretos:")
NUMBERED([
    "No se puede referir a una respuesta específica de EVA. Las respuestas no tienen identidad propia — están embebidas como texto suelto. Por eso es imposible hoy permitir que un usuario haga 'pulgar abajo' sobre una respuesta.",
    "La heurística clasifica mal cuando hay negaciones. La frase 'no me gustó' contiene tanto 'me gustó' (positivo) como 'no me gustó' (negativo), termina empatada y se etiqueta como 'neutro' cuando es claramente negativa.",
    "No hay forma de mejorar la clasificación sin perder el histórico. Si mañana cambiamos la heurística por un modelo NLP, todas las etiquetas viejas pasan a ser dudosas y no hay manera limpia de comparar el viejo contra el nuevo.",
    "El JSON no es analizable. No se pueden hacer queries del tipo '¿qué porcentaje de feedbacks de organización son negativos?' sin escribir código Python a mano cada vez.",
])

H("1.3 Lo que el modelo de datos resuelve", 2)
P("El modelo propuesto está diseñado, desde el primer día, para que:")
BULLETS([
    "Cada mensaje (del invitado y del agente) tenga una identidad estable y única (un UUID).",
    "Cada feedback referencie un mensaje específico por su identidad, no por su texto.",
    "Convivan múltiples clasificadores en paralelo (heurística vieja, NLP nuevo, feedback explícito del invitado) sin que uno pise al otro.",
    "Los analytics se hagan con SQL estándar y devuelvan resultados confiables.",
])

# =============================================================================
# 2. ¿QUÉ ES UN MODELO DE DATOS?
# =============================================================================
H("2. ¿Qué es un modelo de datos?", 1)
P(
    "Un modelo de datos es la decisión, tomada antes de escribir código, sobre qué "
    "información se va a guardar, cómo se va a estructurar y cómo se va a relacionar "
    "entre sí. Piénsese como el plano arquitectónico de un edificio: no es el edificio, "
    "pero define cuántas habitaciones tiene, dónde están las puertas y cómo conectan."
)
P(
    "En este caso, en lugar de habitaciones tenemos tablas, y en lugar de puertas "
    "tenemos relaciones. El modelo decide cosas como: ¿Una sesión y un mensaje son la "
    "misma cosa o dos cosas distintas? ¿Un feedback puede existir sin un mensaje? "
    "¿Dos clasificadores pueden generar dos etiquetas para el mismo mensaje, o solo "
    "puede haber una?"
)
CALLOUT(
    "Un modelo de datos bien pensado evita que dentro de seis meses tengamos que "
    "reescribir media aplicación porque 'no habíamos previsto este caso'. Un modelo "
    "mal pensado obliga a parchar con código para compensar una estructura rígida."
)

# =============================================================================
# 3. CONCEPTOS PREVIOS
# =============================================================================
H("3. Conceptos previos necesarios", 1)
P("Para leer las siguientes secciones se usan seis conceptos. Acá van uno por uno con un ejemplo.")

H("3.1 Tabla, fila y columna", 2)
P(
    "Una base de datos relacional guarda información en tablas, que se parecen a las "
    "hojas de cálculo de Excel. Cada tabla tiene un nombre (ej. 'sessions'), un conjunto "
    "fijo de columnas (cada columna tiene un nombre y un tipo de dato) y un número "
    "variable de filas (cada fila es un registro concreto)."
)
P("Ejemplo — tres filas de la tabla sessions:")
CODE("""id                                  | started_at           | event_id
------------------------------------+----------------------+----------------------
323de5f6-0b4c-40b5-853a-8a39a35e1775 | 2026-05-30T18:00Z    | diplomas-ahk-2026
c2cb64f2-586d-4f6f-bd36-7898a51e18f3 | 2026-05-30T18:15Z    | diplomas-ahk-2026
f440c84b-f4c0-4cf7-86bb-08d61583e0ee | 2026-05-30T18:30Z    | diplomas-ahk-2026""")

H("3.2 Tipo de dato", 2)
P(
    "Cada columna acepta solo un tipo de información. Los tipos básicos que usamos en "
    "este modelo son:"
)
TABLE(
    ["Tipo", "Qué representa", "Ejemplo"],
    [
        ["TEXT", "Cadena de caracteres", "'hola', '323de5f6-...'"],
        ["INTEGER", "Número entero", "1, 42, 1450"],
        ["REAL", "Número con decimales", "0.88, -0.78"],
        ["BLOB", "Bytes crudos (binario)", "el vector de embeddings"],
    ],
    widths=[Inches(1.2), Inches(2.5), Inches(2.7)],
)
CALLOUT(
    "SQLite (la base que usa esta POC) no tiene un tipo nativo para UUID ni para "
    "fecha-hora; usamos TEXT y guardamos los valores en formato estándar (UUID como "
    "string, fechas en ISO-8601). Postgres, en cambio, sí tiene tipos UUID y "
    "TIMESTAMPTZ; cuando migremos solo cambian los tipos, no los datos."
)

H("3.3 Clave primaria (Primary Key)", 2)
P(
    "Es la columna (o combinación de columnas) que identifica de manera única a cada "
    "fila. No puede repetirse y no puede ser nula. Es como el DNI de una persona: dos "
    "personas distintas no pueden tener el mismo DNI."
)
P(
    "En nuestro modelo, la clave primaria de casi todas las tablas es un UUID (un "
    "identificador de 128 bits que se genera aleatoriamente y que prácticamente no se "
    "repite). La excepción es feedback_categories, donde usamos el nombre corto de la "
    "categoría ('ceremonia', 'catering') porque ese nombre ya es único por sí mismo."
)

H("3.4 Clave foránea (Foreign Key)", 2)
P(
    "Es una columna que apunta a la clave primaria de OTRA tabla. Sirve para expresar "
    "que dos filas están relacionadas, y para que la base prohíba estados inconsistentes."
)
P(
    "Ejemplo: la tabla feedback tiene una columna message_id que es FK hacia messages.id. "
    "Eso significa que no se puede insertar un feedback si el mensaje al que apunta no "
    "existe; y si borramos un mensaje, la base puede borrar automáticamente todos sus "
    "feedbacks (ON DELETE CASCADE)."
)

H("3.5 Restricciones (CHECK, UNIQUE, NOT NULL)", 2)
P(
    "Reglas que la base aplica automáticamente cada vez que se intenta insertar o "
    "modificar una fila. Si la regla no se cumple, la operación falla."
)
BULLETS([
    "NOT NULL: la columna no admite valores nulos. Ej. raw_signal en feedback nunca puede ser vacío.",
    "CHECK: una expresión que debe ser verdadera. Ej. rating tiene CHECK (rating BETWEEN 1 AND 5); poner 9 falla.",
    "UNIQUE: no puede haber dos filas con el mismo valor (o combinación). Ej. UNIQUE (session_id, turn_index, role) en messages.",
])
CALLOUT(
    "Las restricciones son la frontera de defensa más barata contra datos basura. "
    "Una restricción en la base le ahorra al desarrollador escribir ese chequeo en "
    "cada lugar del código donde se inserta."
)

H("3.6 Índice", 2)
P(
    "Una estructura auxiliar (similar al índice alfabético de un libro) que permite "
    "buscar filas por una columna mucho más rápido. Sin índice, la base lee todas las "
    "filas para encontrar las que coinciden; con índice, salta directo."
)
P("Ejemplo: si queremos rankear categorías por volumen, este query")
CODE("SELECT category, COUNT(*) FROM feedback GROUP BY category;")
P(
    "se beneficia del índice idx_feedback_category. En esta POC el dataset es chico "
    "y no se nota; en producción con millones de filas, hace la diferencia entre "
    "segundos y minutos."
)

H("3.7 Cardinalidad — 1:1, 1:N, N:M", 2)
P(
    "La cardinalidad describe cuántas filas de una tabla pueden corresponder a cuántas "
    "filas de otra."
)
TABLE(
    ["Relación", "Significado", "Ejemplo en este modelo"],
    [
        ["1:1", "Una fila de A se corresponde con UNA fila de B", "feedback ↔ feedback_embedding (cada feedback tiene como máximo un embedding)"],
        ["1:N", "Una fila de A se corresponde con MUCHAS filas de B", "sessions ↔ messages (una sesión tiene muchos mensajes)"],
        ["N:M", "Muchas a muchas (requiere tabla intermedia)", "No la usamos en este modelo, pero ejemplo típico: estudiantes ↔ materias"],
    ],
    widths=[Inches(0.9), Inches(2.7), Inches(2.9)],
)

# =============================================================================
# 4. LAS 5 ENTIDADES
# =============================================================================
doc.add_page_break()
H("4. Las cinco entidades del modelo", 1)
P(
    "El modelo está formado por cinco tablas. Las explico en orden creciente de "
    "complejidad y dependencia, porque cada una se apoya en las anteriores."
)

# 4.1
H("4.1 feedback_categories — el catálogo", 2)
P(
    "Es una tabla chica, casi inmutable, que enumera las categorías válidas de "
    "feedback. Hoy son siete: ceremonia, organización, recepción, catering, contenido, "
    "agente y general."
)
CODE("""CREATE TABLE feedback_categories (
    slug         TEXT PRIMARY KEY,
    label        TEXT NOT NULL,
    description  TEXT,
    active       INTEGER NOT NULL DEFAULT 1
);""")
P("Por qué existe como tabla separada y no como una lista hardcodeada en el código:")
BULLETS([
    "Si mañana queremos agregar una categoría ('música', por ejemplo), basta con INSERT — no se toca código.",
    "Si queremos archivar una categoría (active = 0), las viejas no se borran pero dejan de mostrarse en formularios.",
    "Permite que la tabla feedback valide por clave foránea: si alguien intenta insertar feedback con category='xyz' y 'xyz' no está en el catálogo, la base rechaza la operación.",
])

# 4.2
H("4.2 sessions — la conversación", 2)
P(
    "Una sesión representa una conversación entera entre un invitado y EVA, desde que "
    "el invitado empieza a hablar hasta que termina (o pasa cierto tiempo de "
    "inactividad)."
)
CODE("""CREATE TABLE sessions (
    id              TEXT PRIMARY KEY,            -- UUID
    started_at      TEXT NOT NULL,               -- ISO-8601 UTC
    ended_at        TEXT,                        -- NULL = en curso
    event_id        TEXT NOT NULL DEFAULT 'diplomas-ahk-2026',
    client_metadata TEXT                         -- JSON: locale, device, etc.
);""")
P("Decisiones:")
BULLETS([
    "id es UUID — no autoincremental — para que el frontend pueda generar IDs sin coordinar con el servidor y para que no haya colisiones si en el futuro mergeamos datos de varios eventos.",
    "ended_at puede ser NULL — sirve para saber qué sesiones están en curso vs cuáles ya terminaron.",
    "event_id permite, el día que el sistema se use en más de un evento, separar los datos sin tocar la estructura.",
    "client_metadata es un campo JSON libre para guardar metadatos que no queremos modelar todavía (qué dispositivo, qué idioma, etc.).",
])

# 4.3
H("4.3 messages — el turno", 2)
P(
    "Cada mensaje individual (sea del invitado o de EVA) es una fila acá. Esta tabla es "
    "la diferencia más importante con el sistema actual de sage-agent: hoy los mensajes "
    "viven en una lista en memoria sin identidad; acá cada uno tiene un UUID estable."
)
CODE("""CREATE TABLE messages (
    id                 TEXT PRIMARY KEY,
    session_id         TEXT NOT NULL REFERENCES sessions(id) ON DELETE CASCADE,
    role               TEXT NOT NULL CHECK (role IN ('user','assistant','system')),
    content            TEXT NOT NULL,
    created_at         TEXT NOT NULL,
    turn_index         INTEGER NOT NULL,
    parent_message_id  TEXT REFERENCES messages(id),
    model_name         TEXT,
    token_count        INTEGER,
    latency_ms         INTEGER,
    UNIQUE (session_id, turn_index, role)
);""")
P("Por qué cada campo está:")
TABLE(
    ["Campo", "Para qué"],
    [
        ["id", "Identidad estable. Sin esto no se puede referir a 'esta respuesta de EVA'."],
        ["session_id", "FK que vincula con sessions. ON DELETE CASCADE = si se borra la sesión, se borran los mensajes."],
        ["role", "Distingue invitado vs agente vs system prompt. CHECK garantiza los 3 valores."],
        ["content", "El texto del mensaje."],
        ["created_at", "Cuándo se produjo. Permite reconstruir la conversación en orden."],
        ["turn_index", "Número de turno dentro de la sesión (1, 2, 3...)."],
        ["parent_message_id", "Apunta del mensaje del asistente al mensaje del usuario que lo motivó. Permite reconstruir pares pregunta-respuesta."],
        ["model_name", "Qué modelo generó la respuesta (NULL si rol = user). Útil para comparar versiones del modelo."],
        ["token_count, latency_ms", "Métricas técnicas opcionales para diagnosticar rendimiento."],
    ],
    widths=[Inches(1.6), Inches(4.6)],
)
CALLOUT(
    "La restricción UNIQUE (session_id, turn_index, role) garantiza que no haya dos "
    "mensajes con el mismo turno y rol en la misma sesión — protege contra duplicados "
    "por reintentos del frontend."
)

# 4.4
H("4.4 feedback — la evaluación", 2)
P(
    "Es la tabla central de todo el sistema. Cada fila representa una evaluación de UN "
    "mensaje (típicamente una respuesta de EVA). Un mensaje puede tener varias "
    "evaluaciones — esa es la decisión de diseño más importante."
)
CODE("""CREATE TABLE feedback (
    id                  TEXT PRIMARY KEY,
    message_id          TEXT NOT NULL REFERENCES messages(id) ON DELETE CASCADE,
    session_id          TEXT NOT NULL REFERENCES sessions(id) ON DELETE CASCADE,
    source              TEXT NOT NULL CHECK (source IN
                          ('auto_heuristic','auto_nlp','user_explicit','ia_team_manual')),
    classifier_version  TEXT NOT NULL,
    sentiment_label     TEXT CHECK (sentiment_label IN
                          ('positive','negative','neutral','mixed')),
    sentiment_score     REAL CHECK (sentiment_score BETWEEN -1.0 AND 1.0),
    category            TEXT REFERENCES feedback_categories(slug),
    subcategory         TEXT,
    rating              INTEGER CHECK (rating BETWEEN 1 AND 5),
    thumbs              TEXT CHECK (thumbs IN ('up','down')),
    comment             TEXT,
    raw_signal          TEXT NOT NULL,
    confidence          REAL CHECK (confidence BETWEEN 0 AND 1),
    created_at          TEXT NOT NULL,
    processed_at        TEXT,
    metadata            TEXT
);""")
P("Campos agrupados por propósito:")
TABLE(
    ["Grupo", "Campos", "Para qué"],
    [
        ["Identidad", "id, message_id, session_id", "Quién es y a qué se refiere."],
        ["Origen", "source, classifier_version", "Quién/qué emitió esta evaluación y con qué versión."],
        ["Clasificación auto", "sentiment_label, sentiment_score, category, subcategory, confidence", "Lo que dijo el clasificador."],
        ["Clasificación humana", "rating, thumbs, comment", "Lo que dijo el invitado (UI) o el equipo IA."],
        ["Auditoría", "raw_signal, created_at, processed_at, metadata", "Texto crudo + timestamps + JSON libre."],
    ],
    widths=[Inches(1.5), Inches(2.4), Inches(2.7)],
)
P(
    "Casi todos los campos son opcionales a propósito. Un feedback puede ser puro "
    "thumbs (sin categoría ni texto), puro comentario (sin rating), o puro automático "
    "(sin thumbs ni comment). El schema permite las tres formas sin obligar a llenar "
    "con NULLs forzados."
)

# 4.5
H("4.5 feedback_embeddings — el vector (futuro NLP)", 2)
P(
    "Cuando se reemplace la heurística por un modelo NLP, cada feedback va a tener "
    "asociado un embedding: un vector numérico de varios cientos de dimensiones que "
    "representa el 'significado' del mensaje en un espacio matemático. Esos vectores "
    "se usan para buscar mensajes parecidos, agrupar feedbacks por temática y entrenar "
    "modelos de clasificación."
)
CODE("""CREATE TABLE feedback_embeddings (
    feedback_id  TEXT PRIMARY KEY REFERENCES feedback(id) ON DELETE CASCADE,
    model_name   TEXT NOT NULL,
    dimensions   INTEGER NOT NULL,
    vector       BLOB NOT NULL,
    created_at   TEXT NOT NULL
);""")
P(
    "Está en una tabla aparte por dos razones: los vectores son grandes (cientos de "
    "floats por fila) y la mayoría de los queries analíticos NO los necesitan. "
    "Separarlos hace los queries comunes mucho más rápidos. En SQLite el vector se "
    "guarda como BLOB; en Postgres con la extensión pgvector sería un tipo nativo "
    "vector(384) que soporta búsqueda por similitud directamente con SQL."
)

# =============================================================================
# 5. DIAGRAMA DE RELACIONES
# =============================================================================
doc.add_page_break()
H("5. Diagrama de relaciones (ER)", 1)
P(
    "Las flechas marcan dirección de la dependencia: el lado donde está la flecha tiene "
    "la clave foránea que apunta al lado opuesto."
)
CODE("""    ┌──────────────────────┐
    │ feedback_categories  │
    │  PK: slug            │
    └──────────▲───────────┘
               │ (FK: category)
               │
┌──────────┐  1│   N  ┌──────────┐  1   N  ┌────────────────────┐  1  0..1  ┌───────────────────────┐
│ sessions ├───┴──────┤ messages ├─────────┤      feedback      ├───────────┤ feedback_embeddings   │
│  PK: id  │ session_ │  PK: id  │ message_│       PK: id        │ feedback_│       PK: feedback_id │
└──────────┘   id     └─────▲────┘   id    └─────────────────────┘    id    └───────────────────────┘
                            │
                            │ parent_message_id (auto-ref)
                            │
                       (assistant → user)

Cardinalidades:
  sessions    1 ──── N    messages          (una sesión, muchos mensajes)
  messages    1 ──── N    feedback          (un mensaje, varios feedbacks — clave!)
  feedback    1 ── 0..1   feedback_embeddings (1:0 o 1:1)
  feedback    N ──── 1    feedback_categories (muchos feedbacks comparten categoría)
  messages    1 ── 0..1   messages           (auto-referencia: parent_message_id)""")

CALLOUT(
    "La relación más importante de leer es messages 1:N feedback. Es lo que permite "
    "que la heurística vieja, el NLP nuevo y el thumbs del usuario coexistan como tres "
    "filas distintas sobre el mismo mensaje, sin pisarse."
)

# =============================================================================
# 6. DECISIONES DE DISEÑO
# =============================================================================
H("6. Decisiones de diseño explicadas", 1)

H("6.1 ¿Por qué un mensaje puede tener varios feedbacks?", 2)
P(
    "Es la decisión central. La alternativa más simple sería: cada mensaje tiene como "
    "máximo un feedback (relación 1:0..1). Esa decisión simplifica los queries pero "
    "rompe contra dos casos reales:"
)
NUMBERED([
    "Cuando se cambie la heurística por NLP, queremos comparar las dos clasificaciones del mismo mensaje para saber cuál es mejor. Si solo guardamos una, perdemos esa capacidad.",
    "Un mismo mensaje puede recibir clasificación automática Y feedback explícito del invitado (thumbs). Son señales distintas con valor distinto — meter las dos en una sola fila las mezcla.",
])
P(
    "Por eso el modelo permite N feedbacks por mensaje, distinguidos por los campos "
    "source y classifier_version."
)

H("6.2 ¿Por qué embeddings en tabla separada?", 2)
P(
    "Los vectores son grandes (un embedding típico ocupa ~3 KB). Si los guardáramos "
    "como columna de la tabla feedback, cualquier SELECT que solo quiera el sentimiento "
    "tendría que leer también el vector. Separándolos, los queries analíticos comunes "
    "tocan solo lo que necesitan y son hasta un orden de magnitud más rápidos."
)
P(
    "Como bonus, separar la tabla nos permite tener MÁS de un embedding por feedback "
    "en el futuro (uno por cada modelo que probemos), simplemente cambiando la PK por "
    "(feedback_id, model_name)."
)

H("6.3 ¿Por qué un catálogo en lugar de un enum?", 2)
P(
    "Podríamos haber dejado category como un campo TEXT con valores válidos definidos "
    "en el código. El problema: agregar o cambiar una categoría requeriría editar "
    "código, redeploy, y todos los lugares donde aparezca la lista (validaciones, "
    "formularios, dashboards) tendrían que actualizarse en sincronía."
)
P(
    "Con un catálogo en la base, la lista de categorías es un dato consultable por "
    "API (GET /feedback/categories). Cualquier cliente — frontend, pipeline IA, "
    "dashboards — la lee dinámicamente y se mantiene sincronizado solo."
)

H("6.4 ¿Por qué SQL y no NoSQL (MongoDB)?", 2)
TABLE(
    ["Criterio", "SQL (Postgres/SQLite)", "NoSQL (MongoDB)"],
    [
        ["Datos relacionales (FK estrictas)", "Sí, nativo", "Manual"],
        ["Analytics con JOIN, GROUP BY", "Excelente", "Verboso (aggregation pipeline)"],
        ["Búsqueda vectorial futura", "pgvector", "Atlas Vector Search (cloud)"],
        ["Volumen esperado (bajo/moderado)", "Sobra", "Sobra"],
        ["Familiaridad típica", "Alta", "Media"],
    ],
    widths=[Inches(2.3), Inches(2.2), Inches(2.2)],
)
P(
    "El caso de uso es altamente relacional (feedback depende de message depende de "
    "session) y los queries más importantes son analíticos (CSAT, breakdown por "
    "categoría). En esos dos ejes SQL gana. La 'flexibilidad de schema' de MongoDB no "
    "se está aprovechando porque el schema sí tiene una forma definida."
)

H("6.5 ¿Por qué SQLite en la POC?", 2)
P(
    "SQLite es una base de datos completa que vive en un único archivo, sin servidor. "
    "Es perfecta para una POC: no requiere instalar nada extra, los datos se inspeccionan "
    "con cualquier cliente, y todo el código SQL que escribimos hoy funciona en Postgres "
    "mañana con cambios mínimos (TEXT → UUID/TIMESTAMPTZ, BLOB → vector). La POC "
    "demuestra el modelo; la decisión de motor productivo se posterga."
)

# =============================================================================
# 7. DEMOSTRACIÓN
# =============================================================================
doc.add_page_break()
H("7. Demostración: queries que prueban el valor", 1)
P(
    "Los siguientes seis queries se ejecutan en demo_queries.py. Cada uno responde una "
    "pregunta de negocio que el sistema actual (JSON plano) no puede responder sin "
    "código ad-hoc."
)

H("7.1 Distribución de sentimientos", 2)
CODE("""SELECT sentiment_label, COUNT(*) AS n
FROM feedback
WHERE sentiment_label IS NOT NULL
GROUP BY sentiment_label
ORDER BY n DESC;""")
P("Respuesta: 'positive' 3, 'negative' 2, 'mixed' 1, 'neutral' 1.")
P(
    "Es el indicador más básico. Sirve como heartbeat: si un día sale 80% negative en "
    "una hora, algo se rompió en el evento o en EVA."
)

H("7.2 Top categorías por volumen", 2)
CODE("""SELECT category, COUNT(*) AS n_feedback
FROM feedback
WHERE category IS NOT NULL
GROUP BY category
ORDER BY n_feedback DESC;""")
P("Permite ver dónde está la conversación más activa: si 'organización' domina, el equipo de logística sabe que ese es el tema del día.")

H("7.3 Acuerdo entre clasificadores (auditoría de calidad)", 2)
CODE("""SELECT a.message_id,
       a.sentiment_label AS heuristic_v1,
       b.sentiment_label AS nlp_v2,
       CASE WHEN a.sentiment_label = b.sentiment_label
            THEN 'agree' ELSE 'DISAGREE' END AS verdict
FROM feedback a
JOIN feedback b
  ON a.message_id = b.message_id
 AND a.source = 'auto_heuristic'
 AND b.source = 'auto_nlp';""")
P("Este query es la justificación práctica de toda la decisión 'múltiples feedbacks por mensaje'. Resultado real de la POC:")
CODE("""heuristic_v1 | nlp_v2    | verdict
-------------|-----------|--------
positive     | positive  | agree
positive     | mixed     | DISAGREE
neutral      | negative  | DISAGREE""")
P(
    "Dos desacuerdos sobre tres casos. El caso 'neutral vs negative' es exactamente el "
    "bug del detector de negaciones — la heurística decía 'neutral' a 'no me gustó la "
    "organización...'. Sin este modelo de datos, no podríamos demostrar la mejora "
    "objetivamente."
)

H("7.4 Respuestas problemáticas de EVA (input para mejorar el prompt)", 2)
CODE("""SELECT m.id AS message_id,
       substr(m.content, 1, 60) AS respuesta_eva,
       f.sentiment_label, f.category, f.source
FROM feedback f
JOIN messages m ON m.id = f.message_id
WHERE f.sentiment_label = 'negative' AND m.role = 'assistant'
ORDER BY f.created_at DESC;""")
P("Devuelve las respuestas concretas de EVA que recibieron feedback negativo. Es input directo para iterar el system prompt — sin este modelo, las respuestas y los feedbacks están desconectados.")

H("7.5 Timeline por sesión", 2)
CODE("""SELECT s.id, s.started_at, COUNT(f.id) AS n_feedback,
       SUM(CASE WHEN f.sentiment_label='positive' THEN 1 ELSE 0 END) AS pos,
       SUM(CASE WHEN f.sentiment_label='negative' THEN 1 ELSE 0 END) AS neg,
       SUM(CASE WHEN f.thumbs='down' THEN 1 ELSE 0 END) AS thumbs_down
FROM sessions s
LEFT JOIN feedback f ON f.session_id = s.id
GROUP BY s.id, s.started_at
ORDER BY s.started_at;""")
P("Vista por sesión: cuántos feedbacks, distribución de sentimiento, cuántos thumbs-down. Es la base para identificar sesiones problemáticas y hacer drill-down a la conversación.")

H("7.6 CSAT filtrando por clasificador", 2)
CODE("""SELECT ROUND(
         1.0 * SUM(CASE WHEN sentiment_label = 'positive' THEN 1 ELSE 0 END)
         / NULLIF(SUM(CASE WHEN sentiment_label IN ('positive','negative')
                            THEN 1 ELSE 0 END), 0),
         3
       ) AS csat_nlp_v2,
       COUNT(*) AS n_total
FROM feedback
WHERE source = 'auto_nlp' AND classifier_version = 'embeddings-v2';""")
P(
    "Calculamos el CSAT (% de positivos sobre los etiquetados) usando solo el "
    "clasificador NLP v2, ignorando la heurística vieja. Esto demuestra por qué "
    "necesitamos guardar source y classifier_version: sin esos campos, mezclaríamos "
    "manzanas con peras."
)

# =============================================================================
# 8. CÓMO CORRER LA POC
# =============================================================================
H("8. Cómo correr la prueba de concepto", 1)
CODE("""# desde la carpeta poc_data_model/

# 1) Crear la base vacía con todas las tablas
python build_db.py --reset

# 2) Poblarla con datos de demostración
python seed_demo.py

# 3) Correr los queries de analytics
python demo_queries.py""")
P(
    "La DB queda en sage_analytics.db. Para inspeccionarla a mano se puede usar "
    "cualquier cliente SQLite (DB Browser for SQLite es gratis y multiplataforma). "
    "El archivo está en .gitignore y no se sube al repo."
)

# =============================================================================
# 9. PRÓXIMOS PASOS
# =============================================================================
H("9. Próximos pasos sugeridos", 1)
NUMBERED([
    "Integrar el modelo con feedback_api.py: cambiar la tabla única actual por las cinco tablas de esta POC.",
    "Cuando el equipo IA tenga su NLP listo, hacer que postee con source='auto_nlp' y un classifier_version propio — coexiste con la heurística sin conflicto.",
    "Agregar al frontend la posibilidad de thumbs/rating sobre cada respuesta de EVA (requiere que /chat devuelva message_id).",
    "Migrar a Postgres cuando el volumen lo justifique: cambian solo los tipos del CREATE TABLE y el driver de conexión.",
    "Activar pgvector y poblar feedback_embeddings para habilitar búsqueda semántica.",
])

# =============================================================================
# 10. GLOSARIO
# =============================================================================
H("10. Glosario rápido para la presentación", 1)
TABLE(
    ["Término", "Definición de 1 línea"],
    [
        ["UUID", "Identificador único universal de 128 bits, generable sin coordinación entre máquinas."],
        ["Clave primaria (PK)", "Columna(s) que identifican unívocamente cada fila de una tabla."],
        ["Clave foránea (FK)", "Columna que referencia la PK de otra tabla; expresa relación."],
        ["Cardinalidad", "Cuántas filas de una tabla se corresponden con cuántas de otra (1:1, 1:N, N:M)."],
        ["Restricción CHECK", "Regla que la base aplica a cada fila para impedir valores inválidos."],
        ["Índice", "Estructura auxiliar que acelera las búsquedas por una columna."],
        ["Embedding", "Vector numérico que representa el significado de un texto en un espacio matemático."],
        ["pgvector", "Extensión de Postgres que agrega tipo vector y búsqueda por similitud."],
        ["CSAT", "Customer Satisfaction Score; típicamente % de respuestas positivas sobre el total clasificable."],
        ["A/B de clasificadores", "Comparar dos clasificadores corriendo en paralelo sobre los mismos datos para elegir el mejor."],
    ],
    widths=[Inches(2.0), Inches(4.5)],
)

# Footer
doc.add_paragraph()
f = doc.add_paragraph(); f.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = f.add_run("— Fin del documento — Prueba de concepto: carpeta poc_data_model/")
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUT)
print(f"OK: {OUT}")
